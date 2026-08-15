"""Editable New Wave IT operational document templates and review exports."""

from __future__ import annotations

import shutil
import tempfile
from html import escape
from pathlib import Path

import fitz
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table as WordTable
from docx.text.paragraph import Paragraph as WordParagraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw
from reportlab.lib import colors as reportlab_colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from brandkit.spec import BRAND, COLORS

DOCUMENT_FILENAMES = (
    "new-wave-it-letterhead.docx",
    "new-wave-it-proposal-template.docx",
    "new-wave-it-capabilities-sheet.docx",
    "new-wave-it-how-to-reach-us.docx",
    "new-wave-it-cybersecurity-notice.docx",
    "new-wave-it-meeting-notes.docx",
    "new-wave-it-project-status.docx",
)

_ROOT = "08-documents-presentations"
_SITE = "newwaveit.com"
_DEEP = COLORS["deep_current"]
_NAVY = COLORS["current_navy"]
_CYAN = COLORS["signal_cyan"]
_GREEN = COLORS["continuity_green"]
_MIST = COLORS["mist_gray"]
_SLATE = COLORS["slate"]
_WHITE = COLORS["pure_white"]


def _rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color.removeprefix("#"))


def _set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill.removeprefix("#"))
    properties.append(shading)


def _set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_widths(table, widths: tuple[float, ...]) -> None:
    table.autofit = False
    grid = table._tbl.tblGrid.gridCol_lst
    for column, width in zip(grid, widths, strict=True):
        column.set(qn("w:w"), str(int(width * 1440)))
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            cell.width = Inches(width)
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _set_table_borders(table, color: str, *, size: int = 4, hide_inside: bool = False) -> None:
    """Write direct OOXML borders so Word never falls back to automatic black."""
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "nil" if hide_inside and edge_name.startswith("inside") else "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color.removeprefix("#"))


def _style(document: Document, name: str, font: str, size: float, color: str, *, bold: bool = False, base: str | None = None, before: float = 0, after: float = 6) -> None:
    styles = document.styles
    style = styles[name] if name in styles else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base:
        style.base_style = styles[base]
    style.font.name = font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = _rgb(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.12


def _apply_theme(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    _style(document, "Normal", "Inter", 10.5, _NAVY, after=7)
    _style(document, "Title", "Plus Jakarta Sans", 28, _DEEP, bold=True, after=8)
    _style(document, "Subtitle", "Inter", 12, _SLATE, after=16)
    _style(document, "Heading 1", "Plus Jakarta Sans", 16, _NAVY, bold=True, before=16, after=6)
    _style(document, "Heading 2", "Plus Jakarta Sans", 12, _NAVY, bold=True, before=12, after=4)
    _style(document, "Heading 3", "Inter", 10.5, _NAVY, bold=True, before=8, after=3)
    _style(document, "Body", "Inter", 10.5, _NAVY, base="Normal", after=7)
    _style(document, "Small Label", "Inter", 8.5, _SLATE, bold=True, after=2)
    _style(document, "Technical Label", "IBM Plex Mono", 8.5, _SLATE, bold=True, after=2)
    _style(document, "Callout", "Inter", 10.5, _NAVY, bold=False, after=8)
    _style(document, "Table Header", "Inter", 9, _WHITE, bold=True, after=0)
    _style(document, "Footer", "Inter", 8, _SLATE, after=0)
    for name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        document.styles[name].paragraph_format.keep_with_next = True


def _logo_path(root: Path) -> Path:
    candidates = (
        root / "02-logos/png/new-wave-it-horizontal-full-color-1024.png",
        root / "02-logos/png/new-wave-it-horizontal-full-color-512.png",
    )
    for candidate in candidates:
        if candidate.is_file():
            return candidate
    raise FileNotFoundError("Task 8 requires the canonical raster logo from Task 4.")


def _add_header_footer(document: Document, root: Path) -> None:
    section = document.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    _set_table_widths(table, (3.6, 2.9))
    left, right = table.rows[0].cells
    left.paragraphs[0].add_run().add_picture(str(_logo_path(root)), width=Inches(1.55))
    right_paragraph = right.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_paragraph.style = document.styles["Technical Label"]
    right_paragraph.add_run("FORT LAUDERDALE | SOUTH FLORIDA\n")
    right_paragraph.add_run(_SITE)
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.style = document.styles["Footer"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(f"{BRAND['name']}  |  {_SITE}  |  {BRAND['support_email']}")


def _document(root: Path) -> Document:
    document = Document()
    _apply_theme(document)
    _add_header_footer(document, root)
    return document


def _add_title(document: Document, title: str, subtitle: str) -> None:
    document.add_paragraph(title, style="Title")
    document.add_paragraph(subtitle, style="Subtitle")


def _add_label(document: Document, text: str, *, technical: bool = False) -> None:
    document.add_paragraph(text, style="Technical Label" if technical else "Small Label")


def _add_callout(document: Document, label: str, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    _set_table_widths(table, (6.5,))
    _set_table_borders(table, _CYAN, size=6, hide_inside=True)
    cell = table.cell(0, 0)
    _set_cell_shading(cell, _MIST)
    paragraph = cell.paragraphs[0]
    paragraph.style = document.styles["Callout"]
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(text)
    document.add_paragraph("", style="Body")


def _add_table(document: Document, headers: tuple[str, ...], rows: tuple[tuple[str, ...], ...], widths: tuple[float, ...]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    _set_table_widths(table, widths)
    _set_table_borders(table, _MIST)
    for cell, text in zip(table.rows[0].cells, headers, strict=True):
        _set_cell_shading(cell, _NAVY)
        paragraph = cell.paragraphs[0]
        paragraph.style = document.styles["Table Header"]
        paragraph.add_run(text)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            paragraph = cell.paragraphs[0]
            paragraph.style = document.styles["Body"]
            paragraph.add_run(value)
    document.add_paragraph("", style="Body")


def _save(document: Document, root: Path, filename: str) -> Path:
    target = root / _ROOT / filename
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)
    return target


def build_letterhead(root: Path) -> Path:
    document = _document(root)
    _add_label(document, "CORRESPONDENCE TEMPLATE", technical=True)
    document.add_paragraph("New Wave IT", style="Title")
    document.add_paragraph("Technology that keeps business moving.", style="Subtitle")
    _add_table(document, ("DATE", "TO", "SUBJECT"), (("[Month DD, YYYY]", "[Recipient name / organization]", "[Purpose of this correspondence]"),), (1.35, 2.4, 2.75))
    document.add_paragraph("[Begin correspondence here. This area is intentionally open for an editable letter, update, or follow-up.]", style="Body")
    return _save(document, root, DOCUMENT_FILENAMES[0])


def build_proposal_template(root: Path) -> Path:
    document = _document(root)
    _add_label(document, "PROPOSAL TEMPLATE", technical=True)
    _add_title(document, "A clear path for dependable technology.", "Prepared by New Wave IT | Fort Lauderdale and South Florida")
    _add_table(document, ("CLIENT / ORGANIZATION", "PROJECT NAME", "PREPARED ON"), (("[Client name]", "[Project name]", "[Month DD, YYYY]"),), (2.25, 2.4, 1.85))
    _add_callout(document, "Purpose", "This editable proposal frames the outcome, scope, timing, and investment conversation without assuming services or commitments not yet agreed.")
    document.add_paragraph("Scope and outcomes", style="Heading 1")
    document.add_paragraph("[Describe the current need, the intended outcome, and the practical work to be performed.]", style="Body")
    document.add_paragraph("Timeline", style="Heading 1")
    _add_table(document, ("PHASE", "TARGET WINDOW", "WORKING OUTPUT"), (("Discover", "[Dates]", "[Shared understanding of needs]"), ("Deliver", "[Dates]", "[Configured or documented work]"), ("Review", "[Dates]", "[Next decisions and handoff]")), (1.35, 1.65, 3.5))
    document.add_paragraph("Investment", style="Heading 1")
    _add_table(document, ("ITEM", "AMOUNT", "NOTES"), (("[Service or project item]", "[$ amount / TBD]", "[Assumptions or approval notes]"),), (2.8, 1.3, 2.4))
    document.add_paragraph("Next steps", style="Heading 1")
    document.add_paragraph("[List the next decision, owner, and agreed follow-up.]")
    return _save(document, root, DOCUMENT_FILENAMES[1])


def build_capabilities_sheet(root: Path) -> Path:
    document = _document(root)
    _add_label(document, "CAPABILITIES SHEET", technical=True)
    _add_title(document, "Technology that keeps business moving.", "Practical support for organizations in Fort Lauderdale and South Florida.")
    _add_callout(document, "How we work", "New Wave IT helps teams organize technology decisions, respond to everyday needs, and maintain a clear line of support.")
    for heading, body in (
        ("Managed IT", "Planning, device and account administration, and day-to-day technology coordination shaped around your environment."),
        ("Cybersecurity", "Practical guidance for reducing avoidable risk, strengthening routines, and responding thoughtfully when something looks wrong."),
        ("Cloud and collaboration", "Support for cloud tools, access, and collaboration practices so work remains connected and understandable."),
        ("Responsive support", f"A shared support inbox at {BRAND['support_email']} for clear requests, updates, and coordinated follow-through."),
    ):
        document.add_paragraph(heading, style="Heading 1")
        document.add_paragraph(body, style="Body")
    document.add_paragraph(f"Start a conversation at {_SITE}.", style="Callout")
    return _save(document, root, DOCUMENT_FILENAMES[2])


def build_reach_us_sheet(root: Path) -> Path:
    document = _document(root)
    _add_label(document, "SUPPORT CONTACT GUIDE", technical=True)
    _add_title(document, "How to reach New Wave IT", "Clear information helps the shared support inbox route your request well.")
    document.add_paragraph(f"Send support requests to {BRAND['support_email']}. This shared support inbox keeps the right people informed and helps preserve context across updates.", style="Body")
    document.add_paragraph("Suggested subject line", style="Heading 1")
    document.add_paragraph("[Organization] | [Issue category] | [Short description]", style="Callout")
    document.add_paragraph("Issue categories", style="Heading 1")
    _add_table(document, ("CATEGORY", "USE IT FOR"), (("Access", "Sign-in, account, or permission questions"), ("Device or software", "A workstation, mobile device, application, or update issue"), ("Connectivity", "Internet, Wi-Fi, remote access, or collaboration trouble"), ("Security concern", "A suspicious message, unexpected activity, or possible incident")), (1.6, 4.9))
    document.add_paragraph("Include this information", style="Heading 1")
    document.add_paragraph("Your name and organization, the best way to reach you, what happened, when it started, the device or service involved, and screenshots or error text when safe to share.", style="Body")
    _add_callout(document, "Protect your account", "Do not email passwords, MFA codes, recovery codes, or sensitive credentials.")
    document.add_paragraph(f"For general information, visit {_SITE}. New Wave IT supports organizations in Fort Lauderdale and South Florida.", style="Body")
    return _save(document, root, DOCUMENT_FILENAMES[3])


def build_security_notice(root: Path) -> Path:
    document = _document(root)
    _add_label(document, "INCIDENT GUIDANCE", technical=True)
    _add_title(document, "Cybersecurity notice and response record", "A calm, editable guide for employees and clients when something does not look right.")
    _add_callout(document, "First response", "Pause, preserve what you observed, and contact the shared support inbox. Do not forward suspicious content broadly or attempt to investigate beyond your role.")
    _add_table(document, ("REPORTED BY", "DATE / TIME", "SYSTEM OR ACCOUNT"), (("[Name]", "[Month DD, YYYY / time]", "[Device, service, or account]"),), (2.1, 2.0, 2.4))
    document.add_paragraph("What happened?", style="Heading 1")
    document.add_paragraph("[Describe the message, activity, or concern. Include observable facts and any relevant screenshots or error text.]", style="Body")
    document.add_paragraph("Immediate actions", style="Heading 1")
    _add_table(document, ("ACTION", "OWNER", "STATUS / TIME"), (("Preserve relevant details", "[Name]", "[Open / complete]"), (f"Notify {BRAND['support_email']}", "[Name]", "[Open / complete]"), ("Document business impact or interruption", "[Name]", "[Open / complete]")), (3.1, 1.25, 2.15))
    _add_callout(document, "Credential safety", "Do not send passwords, MFA codes, recovery codes, or sensitive credentials by email.")
    document.add_paragraph("This template supports a measured response record; it does not imply a confirmed breach or a guaranteed outcome.", style="Body")
    return _save(document, root, DOCUMENT_FILENAMES[4])


def build_meeting_notes(root: Path) -> Path:
    document = _document(root)
    _add_label(document, "MEETING NOTES", technical=True)
    _add_title(document, "Meeting notes", "A simple editable record for decisions, actions, and open questions.")
    _add_table(document, ("MEETING", "DATE", "FACILITATOR"), (("[Topic / project]", "[Month DD, YYYY]", "[Name]"),), (2.8, 1.7, 2.0))
    document.add_paragraph("Attendees", style="Heading 1")
    document.add_paragraph("[Names and roles]", style="Body")
    document.add_paragraph("Discussion notes", style="Heading 1")
    document.add_paragraph("[Capture context, options considered, and the useful details behind decisions.]", style="Body")
    document.add_paragraph("Decisions", style="Heading 1")
    _add_table(document, ("DECISION", "OWNER", "RATIONALE / FOLLOW-UP"), (("[Decision]", "[Name]", "[Why it was made / what changes next]"),), (2.25, 1.15, 3.1))
    document.add_paragraph("Actions", style="Heading 1")
    _add_table(document, ("ACTION", "OWNER", "DUE", "STATUS"), (("[Action item]", "[Name]", "[Date]", "[Not started / in progress / complete]"), ("[Action item]", "[Name]", "[Date]", "[Not started / in progress / complete]")), (2.5, 1.0, 0.8, 2.2))
    document.add_paragraph("Open questions", style="Heading 1")
    document.add_paragraph("[Question, owner, and next checkpoint]", style="Body")
    return _save(document, root, DOCUMENT_FILENAMES[5])


def build_project_status(root: Path) -> Path:
    document = _document(root)
    _add_label(document, "PROJECT STATUS", technical=True)
    _add_title(document, "Project status update", "A practical shared view of progress, decisions, risks, and next work.")
    _add_table(document, ("PROJECT", "REPORTING PERIOD", "OVERALL STATUS"), (("[Project name]", "[Dates]", "[On track / watch / needs decision]"),), (2.35, 1.85, 2.3))
    document.add_paragraph("Summary", style="Heading 1")
    document.add_paragraph("[What changed this period, what is working, and what needs attention.]", style="Body")
    document.add_paragraph("Workstream status", style="Heading 1")
    _add_table(document, ("WORKSTREAM", "OWNER", "STATUS", "NEXT MILESTONE"), (("[Workstream]", "[Name]", "[On track / watch / blocked]", "[Date / decision]"), ("[Workstream]", "[Name]", "[On track / watch / blocked]", "[Date / decision]")), (2.1, 1.0, 1.4, 2.0))
    document.add_paragraph("Risks and decisions", style="Heading 1")
    _add_table(document, ("ITEM", "IMPACT", "OWNER", "NEXT ACTION"), (("[Risk, dependency, or decision]", "[What it affects]", "[Name]", "[Action / date]"),), (2.35, 1.45, 0.9, 1.8))
    document.add_paragraph("Next reporting period", style="Heading 1")
    document.add_paragraph("[Planned work, expected decisions, and the next update date.]", style="Body")
    return _save(document, root, DOCUMENT_FILENAMES[6])


def _rl_color(hex_color: str):
    return reportlab_colors.HexColor(hex_color)


def _pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()["BodyText"]
    return {
        "body": ParagraphStyle("NewWaveBody", parent=base, fontName="Helvetica", fontSize=9.5, leading=13, textColor=_rl_color(_NAVY), spaceAfter=7),
        "title": ParagraphStyle("NewWaveTitle", parent=base, fontName="Helvetica-Bold", fontSize=23, leading=28, textColor=_rl_color(_DEEP), spaceAfter=7),
        "subtitle": ParagraphStyle("NewWaveSubtitle", parent=base, fontName="Helvetica", fontSize=11, leading=15, textColor=_rl_color(_SLATE), spaceAfter=14),
        "heading": ParagraphStyle("NewWaveHeading", parent=base, fontName="Helvetica-Bold", fontSize=14, leading=18, textColor=_rl_color(_NAVY), spaceBefore=11, spaceAfter=5),
        "label": ParagraphStyle("NewWaveLabel", parent=base, fontName="Helvetica-Bold", fontSize=8, leading=10, textColor=_rl_color(_SLATE), spaceAfter=3),
        "footer": ParagraphStyle("NewWaveFooter", parent=base, fontName="Helvetica", fontSize=7.5, leading=9, textColor=_rl_color(_SLATE), alignment=TA_CENTER),
        "table": ParagraphStyle("NewWaveTable", parent=base, fontName="Helvetica", fontSize=8.3, leading=10.5, textColor=_rl_color(_NAVY)),
        "table_header": ParagraphStyle("NewWaveTableHeader", parent=base, fontName="Helvetica-Bold", fontSize=7.7, leading=9, textColor=_rl_color(_WHITE)),
    }


def _paragraph_flowable(paragraph, styles: dict[str, ParagraphStyle]):
    text = escape(paragraph.text).replace("\n", "<br/>")
    if not text:
        return Spacer(1, 5)
    name = paragraph.style.name
    if name == "Title":
        style = styles["title"]
    elif name == "Subtitle":
        style = styles["subtitle"]
    elif name.startswith("Heading"):
        style = styles["heading"]
    elif name in {"Small Label", "Technical Label"}:
        style = styles["label"]
    else:
        style = styles["body"]
    return Paragraph(text, style)


def _table_widths(table, total_width: float) -> list[float]:
    grid = table._tbl.tblGrid.gridCol_lst
    values = [int(column.get(qn("w:w")) or 1) for column in grid]
    if not values:
        return [total_width / len(table.columns)] * len(table.columns)
    scale = total_width / sum(values)
    return [value * scale for value in values]


def _table_flowable(table, styles: dict[str, ParagraphStyle]):
    header_cells = table.rows[0].cells
    is_callout = len(header_cells) == 1 and header_cells[0].paragraphs[0].style.name == "Callout"
    data = []
    for row_index, row in enumerate(table.rows):
        style = styles["body"] if is_callout else styles["table_header"] if row_index == 0 else styles["table"]
        data.append([Paragraph(escape(cell.text).replace("\n", "<br/>"), style) for cell in row.cells])
    rendered = Table(data, colWidths=_table_widths(table, 6.5 * inch), repeatRows=1, hAlign="LEFT")
    commands = [
        ("BACKGROUND", (0, 0), (-1, 0), _rl_color(_NAVY)),
        ("TEXTCOLOR", (0, 0), (-1, 0), _rl_color(_WHITE)),
        ("GRID", (0, 0), (-1, -1), 0.35, _rl_color(_MIST)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]
    if is_callout:
        commands = [
            ("BACKGROUND", (0, 0), (-1, -1), _rl_color(_MIST)),
            ("BOX", (0, 0), (-1, -1), 0.6, _rl_color(_CYAN)),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 10),
            ("RIGHTPADDING", (0, 0), (-1, -1), 10),
            ("TOPPADDING", (0, 0), (-1, -1), 8),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ]
    rendered.setStyle(TableStyle(commands))
    return rendered


def _pdf_page_furniture(root: Path):
    logo = _logo_path(root)

    def draw(canvas, document) -> None:
        canvas.saveState()
        canvas.drawImage(str(logo), document.leftMargin, letter[1] - 0.78 * inch, width=1.55 * inch, height=0.39 * inch, preserveAspectRatio=True, mask="auto")
        canvas.setFont("Helvetica-Bold", 7.5)
        canvas.setFillColor(_rl_color(_SLATE))
        canvas.drawRightString(letter[0] - document.rightMargin, letter[1] - 0.52 * inch, "FORT LAUDERDALE | SOUTH FLORIDA")
        canvas.setFont("Helvetica", 7.5)
        canvas.drawRightString(letter[0] - document.rightMargin, letter[1] - 0.66 * inch, _SITE)
        canvas.setStrokeColor(_rl_color(_CYAN))
        canvas.setLineWidth(0.75)
        canvas.line(document.leftMargin, letter[1] - 0.87 * inch, letter[0] - document.rightMargin, letter[1] - 0.87 * inch)
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(_rl_color(_SLATE))
        canvas.drawCentredString(letter[0] / 2, 0.48 * inch, f"{BRAND['name']}  |  {_SITE}  |  {BRAND['support_email']}")
        canvas.restoreState()

    return draw


def _render_review_pdf(root: Path, source_path: Path, target: Path) -> None:
    """Render the generated template structure into a deterministic review PDF."""
    source = Document(source_path)
    styles = _pdf_styles()
    story = []
    story.append(Paragraph("Editable template fields and grids", styles["label"]))
    for child in source.element.body.iterchildren():
        if child.tag == qn("w:p"):
            story.append(_paragraph_flowable(WordParagraph(child, source), styles))
        elif child.tag == qn("w:tbl"):
            story.append(_table_flowable(WordTable(child, source), styles))
            story.append(Spacer(1, 6))
    document = SimpleDocTemplate(str(target), pagesize=letter, leftMargin=0.75 * inch, rightMargin=0.75 * inch, topMargin=1.12 * inch, bottomMargin=0.72 * inch, title=source_path.stem)
    furniture = _pdf_page_furniture(root)
    document.build(story, onFirstPage=furniture, onLaterPages=furniture)


def _export_review_pdfs(document_root: Path) -> tuple[Path, ...]:
    """Create deterministic ReportLab visual-QA PDFs from the template content."""
    package_root = document_root.parent
    exports = []
    for filename in DOCUMENT_FILENAMES:
        source = document_root / filename
        target = document_root / filename.replace(".docx", "-review.pdf")
        _render_review_pdf(package_root, source, target)
        exports.append(target)
    result = tuple(exports)
    if any(not path.is_file() or path.stat().st_size <= 1_000 for path in result):
        raise RuntimeError("Deterministic review renderer did not produce every required PDF.")
    return result


def _make_contact_sheet(root: Path, review_pdfs: tuple[Path, ...]) -> None:
    temporary = Path(tempfile.mkdtemp(prefix="new-wave-it-document-review-"))
    pages: list[tuple[str, Image.Image]] = []
    try:
        for pdf in review_pdfs:
            rendered = fitz.open(pdf)
            try:
                for index, page in enumerate(rendered, start=1):
                    pixmap = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
                    image = Image.frombytes("RGB", (pixmap.width, pixmap.height), pixmap.samples)
                    qa_path = temporary / f"{pdf.stem}-page-{index}.png"
                    image.save(qa_path)
                    if not qa_path.is_file() or qa_path.stat().st_size <= 1_000:
                        raise RuntimeError(f"Could not rasterize {pdf.name} page {index}.")
                    pages.append((f"{pdf.stem.replace('-review', '')} | page {index}", image))
            finally:
                rendered.close()
        if not pages:
            raise RuntimeError("No document review PDF pages were rasterized.")
        tile_width, tile_height = 360, 500
        columns = 3
        rows = (len(pages) + columns - 1) // columns
        sheet = Image.new("RGB", (columns * tile_width, rows * (tile_height + 34)), _MIST)
        draw = ImageDraw.Draw(sheet)
        for index, (label, image) in enumerate(pages):
            x, y = (index % columns) * tile_width, (index // columns) * (tile_height + 34)
            preview = image.copy()
            preview.thumbnail((tile_width - 20, tile_height - 20))
            sheet.paste(preview, (x + (tile_width - preview.width) // 2, y + 10))
            draw.text((x + 10, y + tile_height + 8), label, fill=_NAVY)
        target = root / "12-preview/document-review-contact-sheet.png"
        target.parent.mkdir(parents=True, exist_ok=True)
        sheet.save(target)
    finally:
        shutil.rmtree(temporary, ignore_errors=True)


def export_documents(root: Path) -> tuple[Path, ...]:
    """Generate editable templates and deterministic ReportLab review PDFs."""
    builders = (
        build_letterhead,
        build_proposal_template,
        build_capabilities_sheet,
        build_reach_us_sheet,
        build_security_notice,
        build_meeting_notes,
        build_project_status,
    )
    documents = tuple(builder(root) for builder in builders)
    review_pdfs = _export_review_pdfs(root / _ROOT)
    _make_contact_sheet(root, review_pdfs)
    return documents
