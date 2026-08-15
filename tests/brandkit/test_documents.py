from __future__ import annotations

import inspect
import re
import zipfile
from pathlib import Path

import fitz
import pytest
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches
from lxml import etree
from PIL import Image

from brandkit.cli import build
from brandkit.documents import DOCUMENT_FILENAMES, _export_review_pdfs, export_documents
from brandkit.email_assets import EMAIL_SUPPORT_FILENAMES, export_email_support_assets
from brandkit.logo import export_logo_family
from brandkit.print_assets import PRINT_FILENAMES, export_print_assets
from brandkit.render import export_raster_family
from brandkit.spec import COLORS


@pytest.fixture(scope="module")
def package_root(tmp_path_factory: pytest.TempPathFactory) -> Path:
    root = build(tmp_path_factory.mktemp("documents") / "new-wave-it-brand-package")
    export_logo_family(root)
    export_raster_family(root)
    export_documents(root)
    export_print_assets(root)
    export_email_support_assets(root)
    return root


def test_document_exports_match_the_contract_and_have_review_pdf_pages(package_root: Path) -> None:
    document_root = package_root / "08-documents-presentations"
    expected = set(DOCUMENT_FILENAMES)
    assert {path.name for path in document_root.glob("*.docx")} == expected
    for filename in DOCUMENT_FILENAMES:
        review_pdf = document_root / filename.replace(".docx", "-review.pdf")
        assert review_pdf.is_file() and review_pdf.stat().st_size > 1_000
        rendered = fitz.open(review_pdf)
        try:
            assert rendered.page_count >= 1
        finally:
            rendered.close()
    contact_sheet = package_root / "12-preview/document-review-contact-sheet.png"
    assert contact_sheet.is_file() and contact_sheet.stat().st_size > 1_000


def test_review_pdfs_contain_template_specific_content_and_do_not_use_process_automation(package_root: Path) -> None:
    expected_text = {
        "new-wave-it-letterhead.docx": "CORRESPONDENCE TEMPLATE",
        "new-wave-it-proposal-template.docx": "A clear path for dependable technology.",
        "new-wave-it-capabilities-sheet.docx": "Managed IT",
        "new-wave-it-how-to-reach-us.docx": "Do not email passwords, MFA codes, recovery codes, or sensitive credentials.",
        "new-wave-it-cybersecurity-notice.docx": "This template supports a measured response record",
        "new-wave-it-meeting-notes.docx": "Meeting notes",
        "new-wave-it-project-status.docx": "Project status update",
    }
    for filename, expected in expected_text.items():
        review_pdf = package_root / "08-documents-presentations" / filename.replace(".docx", "-review.pdf")
        rendered = fitz.open(review_pdf)
        try:
            text = "\n".join(page.get_text() for page in rendered)
        finally:
            rendered.close()
        assert expected in text

    source = inspect.getsource(_export_review_pdfs).lower()
    assert all(token not in source for token in ("subprocess", "powershell", "winword", "comobject", "word.application"))


def test_review_pdf_pages_have_shared_furniture_and_contact_sheet_covers_every_page(package_root: Path) -> None:
    document_root = package_root / "08-documents-presentations"
    page_count = 0
    for filename in DOCUMENT_FILENAMES:
        review_pdf = document_root / filename.replace(".docx", "-review.pdf")
        rendered = fitz.open(review_pdf)
        try:
            for page in rendered:
                page_count += 1
                extracted = page.get_text()
                assert "FORT LAUDERDALE | SOUTH FLORIDA" in extracted
                assert "newwaveit.com" in extracted
                assert "support@newwaveitfl.com" in extracted
        finally:
            rendered.close()

    contact_sheet = package_root / "12-preview/document-review-contact-sheet.png"
    with Image.open(contact_sheet) as image:
        columns, tile_width, tile_height, label_height = 3, 360, 500, 34
        rows = (page_count + columns - 1) // columns
        assert image.size == (columns * tile_width, rows * (tile_height + label_height))
        background = (232, 239, 241)
        for page_index in range(page_count):
            x = (page_index % columns) * tile_width
            y = (page_index // columns) * (tile_height + label_height)
            page_tile = image.crop((x + 10, y + 10, x + tile_width - 10, y + tile_height - 10)).convert("RGB")
            label_tile = image.crop((x + 10, y + tile_height, x + tile_width - 10, y + tile_height + label_height)).convert("RGB")
            assert any(pixel != background for pixel in page_tile.getdata())
            assert any(pixel != background for pixel in label_tile.getdata())


def test_every_document_is_editable_letter_sized_and_uses_the_theme(package_root: Path) -> None:
    document_root = package_root / "08-documents-presentations"
    required_styles = {
        "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "Body",
        "Small Label", "Technical Label", "Callout", "Table Header", "Footer",
    }
    for filename in DOCUMENT_FILENAMES:
        document = Document(document_root / filename)
        section = document.sections[0]
        assert section.page_width == Inches(8.5)
        assert section.page_height == Inches(11)
        assert section.top_margin == Inches(1)
        assert section.bottom_margin == Inches(1)
        assert section.left_margin == Inches(1)
        assert section.right_margin == Inches(1)
        assert required_styles <= {style.name for style in document.styles}
        assert document.paragraphs
        assert document.sections[0].header.paragraphs
        assert document.sections[0].footer.paragraphs
        header_text = "\n".join(
            [paragraph.text for paragraph in document.sections[0].header.paragraphs]
            + [cell.text for table in document.sections[0].header.tables for row in table.rows for cell in row.cells]
        )
        footer_text = "\n".join(
            [paragraph.text for paragraph in document.sections[0].footer.paragraphs]
            + [cell.text for table in document.sections[0].footer.tables for row in table.rows for cell in row.cells]
        )
        assert "FORT LAUDERDALE | SOUTH FLORIDA" in header_text
        assert "newwaveit.com" in header_text
        assert "newwaveit.com" in footer_text
        assert "support@newwaveitfl.com" in footer_text


def test_documents_have_no_external_relationships_embedded_fonts_or_legacy_drawingml_colors(package_root: Path) -> None:
    legacy_colors = {"00B2FF", "00AEEF", "002D62", "003E7E", "0B1F3A"}
    for filename in DOCUMENT_FILENAMES:
        with zipfile.ZipFile(package_root / "08-documents-presentations" / filename) as archive:
            names = archive.namelist()
            assert not any(name.startswith("word/fonts/") for name in names)
            for rel_name in (name for name in names if name.endswith(".rels")):
                root = etree.fromstring(archive.read(rel_name))
                assert not any(rel.get("TargetMode") == "External" for rel in root)
            for xml_name in (name for name in names if name.endswith(".xml")):
                root = etree.fromstring(archive.read(xml_name))
                colors = root.xpath("//*[local-name()='solidFill']/*[local-name()='srgbClr']/@val")
                assert legacy_colors.isdisjoint(set(colors))


def test_document_body_tables_use_direct_approved_border_tokens(package_root: Path) -> None:
    document_root = package_root / "08-documents-presentations"
    operational_edges = ("top", "left", "bottom", "right", "insideH", "insideV")
    approved_border_colors = {"E8EFF1", "31C6CF"}
    for filename in DOCUMENT_FILENAMES:
        with zipfile.ZipFile(document_root / filename) as archive:
            root = etree.fromstring(archive.read("word/document.xml"))
        tables = root.xpath("//*[local-name()='body']/*[local-name()='tbl']")
        assert tables
        for table in tables:
            borders = table.xpath("./*[local-name()='tblPr']/*[local-name()='tblBorders']")
            assert len(borders) == 1
            edges = {etree.QName(edge).localname: edge for edge in borders[0]}
            assert set(operational_edges) <= set(edges)

            is_callout = len(table.xpath("./*[local-name()='tr']")) == 1 and len(table.xpath("./*[local-name()='tr']/*[local-name()='tc']")) == 1
            for edge_name in operational_edges:
                edge = edges[edge_name]
                color = edge.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color")
                value = edge.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                assert color in approved_border_colors
                assert color not in {"auto", "000000"}
                if is_callout and edge_name in {"insideH", "insideV"}:
                    assert value == "nil"
                    assert color == "31C6CF"
                elif is_callout:
                    assert value == "single"
                    assert color == "31C6CF"
                else:
                    assert value == "single"
                    assert color == "E8EFF1"
                    assert edge.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz") in {"4", "6"}

            for cell_border in table.xpath(".//*[local-name()='tcBorders']/*"):
                color = cell_border.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color")
                assert color is None or color in approved_border_colors


def test_reach_us_document_has_real_safety_guidance_and_editable_support_workflow(package_root: Path) -> None:
    document = Document(package_root / "08-documents-presentations/new-wave-it-how-to-reach-us.docx")
    text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    assert "support@newwaveitfl.com" in text
    assert "newwaveit.com" in text
    assert "Do not email passwords, MFA codes, recovery codes, or sensitive credentials." in text
    assert "Suggested subject line" in text
    assert "shared support inbox" in text.lower()
    assert "Fort Lauderdale" in text
    assert document.tables


def test_operational_templates_include_editable_fields_and_real_grid_data(package_root: Path) -> None:
    root = package_root / "08-documents-presentations"
    proposal = Document(root / "new-wave-it-proposal-template.docx")
    proposal_text = "\n".join(
        [paragraph.text for paragraph in proposal.paragraphs]
        + [cell.text for table in proposal.tables for row in table.rows for cell in row.cells]
    )
    assert all(marker in proposal_text.casefold() for marker in ("client / organization", "project name", "investment", "next steps"))

    meeting = Document(root / "new-wave-it-meeting-notes.docx")
    status = Document(root / "new-wave-it-project-status.docx")
    assert len(meeting.tables) >= 2
    assert len(status.tables) >= 2
    assert "owner" in "\n".join(cell.text for table in meeting.tables for row in table.rows for cell in row.cells).casefold()
    assert "status" in "\n".join(cell.text for table in status.tables for row in table.rows for cell in row.cells).casefold()


def test_business_cards_have_bleed_trim_safe_copy_and_print_colorways(package_root: Path) -> None:
    print_root = package_root / "09-print"
    assert {path.name for path in print_root.glob("new-wave-it-business-card-*.*")} == set(PRINT_FILENAMES[:8])
    approved_digital = {value.removeprefix("#").upper() for value in COLORS.values()}
    for orientation, media_size in (("horizontal", (4.0, 2.5)), ("vertical", (2.5, 4.0))):
        for colorway in ("full-color", "print-black"):
            stem = f"new-wave-it-business-card-{orientation}-{colorway}"
            svg_path = print_root / f"{stem}.svg"
            pdf_path = print_root / f"{stem}.pdf"
            root = etree.fromstring(svg_path.read_bytes())
            safe_copy = root.xpath("//*[@data-role='safe-text']")
            crop_marks = root.xpath("//*[@data-role='crop-mark']")
            assert len(safe_copy) >= 4
            assert len(crop_marks) == 8
            trim = root.xpath("//*[@data-role='trim-box']")[0]
            bleed = root.xpath("//*[@data-role='bleed-box']")[0]
            trim_x, trim_y = float(trim.get("x")), float(trim.get("y"))
            trim_width, trim_height = float(trim.get("width")), float(trim.get("height"))
            bleed_x, bleed_y = float(bleed.get("x")), float(bleed.get("y"))
            bleed_width, bleed_height = float(bleed.get("width")), float(bleed.get("height"))
            for text_node in safe_copy:
                x, y = float(text_node.get("x")), float(text_node.get("y"))
                assert trim_x + 9 <= x <= trim_x + trim_width - 9
                assert trim_y + 9 <= y <= trim_y + trim_height - 9
            assert any(
                float(mark.get(attribute)) < bleed_x
                or float(mark.get(attribute)) > bleed_x + bleed_width
                or float(mark.get(attribute)) < bleed_y
                or float(mark.get(attribute)) > bleed_y + bleed_height
                for mark in crop_marks
                for attribute in ("x1", "x2", "y1", "y2")
            )
            colors = {value.upper() for value in re.findall(r"#[0-9A-Fa-f]{6}", svg_path.read_text(encoding="utf-8"))}
            colors = {value.removeprefix("#") for value in colors}
            if colorway == "full-color":
                assert colors <= approved_digital
                assert "000000" not in colors
            else:
                assert colors <= {"000000", "FFFFFF"}
            if orientation == "horizontal":
                contrast_panel = root.xpath("//*[@data-role='contrast-panel']")
                assert len(contrast_panel) == 1
                assert float(contrast_panel[0].get("width")) >= 130
            if colorway == "print-black":
                identity_text = [node for node in safe_copy if node.text in {"JORDAN LEE", "Client Success Lead"}]
                assert identity_text
                assert {node.get("fill").upper() for node in identity_text} == {"#FFFFFF"}

            pdf = fitz.open(pdf_path)
            try:
                page = pdf[0]
                assert page.rect.width / 72 == pytest.approx(media_size[0])
                assert page.rect.height / 72 == pytest.approx(media_size[1])
                assert page.trimbox.width / 72 == pytest.approx(3.5 if orientation == "horizontal" else 2.0)
                assert page.trimbox.height / 72 == pytest.approx(2.0 if orientation == "horizontal" else 3.5)
                assert page.bleedbox.width / 72 == pytest.approx(3.75 if orientation == "horizontal" else 2.25)
                assert page.bleedbox.height / 72 == pytest.approx(2.25 if orientation == "horizontal" else 3.75)
                text = page.get_text()
                assert "jordan lee" in text.casefold()
                assert "support@newwaveitfl.com" in text
                if colorway == "full-color":
                    render = page.get_pixmap(alpha=False)
                    sample_x = int(render.width * (0.80 if orientation == "horizontal" else 0.50))
                    sample_y = int(render.height * (0.50 if orientation == "horizontal" else 0.72))
                    pixel = Image.frombytes("RGB", (render.width, render.height), render.samples).getpixel((sample_x, sample_y))
                    assert min(pixel) >= 215
                else:
                    render = page.get_pixmap(matrix=fitz.Matrix(4, 4), alpha=False)
                    image = Image.frombytes("RGB", (render.width, render.height), render.samples)
                    if orientation == "horizontal":
                        identity_region = image.crop((80, 180, 500, 520))
                    else:
                        identity_region = image.crop((80, 60, 500, 460))
                    assert sum(1 for pixel in identity_region.getdata() if min(pixel) >= 235) > 120
            finally:
                pdf.close()


def test_invoice_assets_are_local_raster_vector_pairs_with_usage_pdf(package_root: Path) -> None:
    print_root = package_root / "09-print"
    expected_pairs = {
        "new-wave-it-invoice-header": (1200, 220),
        "new-wave-it-invoice-footer": (1200, 120),
    }
    for stem, dimensions in expected_pairs.items():
        svg = print_root / f"{stem}.svg"
        png = print_root / f"{stem}.png"
        assert svg.is_file() and png.is_file()
        assert "newwaveit.com" in svg.read_text(encoding="utf-8")
        with Image.open(png) as image:
            assert image.size == dimensions
    usage = fitz.open(print_root / "new-wave-it-invoice-asset-usage.pdf")
    try:
        assert usage.page_count == 1
        text = usage[0].get_text()
        assert "Invoice asset usage" in text
        assert "minimum padding" in text.casefold()
        assert "monochrome" in text.casefold()
        assert "Support contact: support@newwaveitfl.com | newwaveit.com" in text
        assert "| newwaveitfl.com" not in text
    finally:
        usage.close()


def test_email_and_support_assets_are_self_contained_and_operational(package_root: Path) -> None:
    root = package_root / "10-email-support"
    assert {path.name for path in root.iterdir()} == set(EMAIL_SUPPORT_FILENAMES)
    signature = (root / "new-wave-it-email-signature.html").read_text(encoding="utf-8")
    assert "<table" in signature.casefold()
    assert "<style" not in signature.casefold()
    assert "new-wave-it-email-signature.png" in signature
    assert "src=\"http" not in signature.casefold()
    assert "support@newwaveitfl.com" in signature
    assert "width:100%;max-width:640px" in signature
    assert "width:100%;max-width:480px" in signature
    assert "display:block;width:480px" not in signature
    assert "Plus Jakarta Sans" in signature
    assert "Inter" in signature
    plain_text = (root / "new-wave-it-email-signature.txt").read_text(encoding="utf-8")
    assert "Jordan Lee" in plain_text
    assert "newwaveit.com" in plain_text
    installation_guide = (root / "new-wave-it-email-signature-installation-guide.md").read_text(encoding="utf-8")
    for client in ("Outlook desktop", "Outlook web", "Gmail", "Apple Mail", "mobile"):
        assert client.casefold() in installation_guide.casefold()

    for filename in EMAIL_SUPPORT_FILENAMES:
        path = root / filename
        if path.suffix == ".html":
            html = path.read_text(encoding="utf-8")
            assert "<table" in html.casefold()
            assert "style=" in html.casefold()
            assert "src=\"http" not in html.casefold()
        elif path.suffix == ".png":
            with Image.open(path) as image:
                assert image.width >= 600 and image.height >= 80

    required_fields = ("status", "impact", "affected service", "started", "next update", "support contact")
    for filename in ("new-wave-it-maintenance-notice.html", "new-wave-it-incident-update.html"):
        operational = (root / filename).read_text(encoding="utf-8").casefold()
        for label in required_fields:
            assert label in operational
        assert "support@newwaveitfl.com" in operational
